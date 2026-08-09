from __future__ import annotations

"""Build a normalized image-only DOCX fixture from source-supported Gerloff SI rows.

This fixture exists only to exercise the multimodal ingestion path in CI. The rows are
source evidence from the supplementary inventory figures; they are rendered as images
so the extraction model cannot obtain them from machine-readable DOCX table text.
The original blind benchmark result remains the historical baseline.
"""

import io
from pathlib import Path

import fitz
from docx import Document


TABLES = {
    "AEC 1 MW supplementary inventory": [
        ("low-alloyed steel for container", "6075.6", "kg"),
        ("concrete for foundation", "7.7", "m3"),
        ("chromium steel for anode and cathode frame", "20194.4", "kg"),
        ("nickel for anode and cathode frame", "2884.9", "kg"),
        ("tetrafluoroethylene for gasket", "144.2", "kg"),
        ("polysulfone for Zirfon diaphragm", "48.8", "kg"),
        ("zirconium oxide for Zirfon diaphragm", "73.0", "kg"),
    ],
    "PEMEC 1 MW supplementary inventory": [
        ("low-alloyed steel for container", "2250.0", "kg"),
        ("concrete for foundation", "2.3", "m3"),
        ("titanium for bipolar plate", "528.0", "kg"),
        ("tetrafluoroethylene for membrane polymer", "16.0", "kg"),
        ("carbon black for electrocatalyst anode", "4.5", "kg"),
        ("iridium for electrocatalyst anode", "0.8", "kg"),
        ("platinum for electrocatalyst cathode", "0.075", "kg"),
        ("copper for current collector", "4.5", "kg"),
        ("synthetic rubber for gasket", "4.8", "kg"),
    ],
    "SOEC 1 MW supplementary inventory": [
        ("low-alloyed steel for container", "2250.0", "kg"),
        ("concrete for foundation", "2.3", "m3"),
        ("chromium steel for air electrode", "8976.1", "kg"),
        ("praseodymium oxide for air electrode screen printing", "9.0", "kg"),
        ("nickel for air electrode screen printing", "7.5", "kg"),
        ("zirconium oxide for blocking layer", "170.7", "kg"),
        ("samarium europium gadolinium concentrate for blocking layer", "14.8", "kg"),
        ("cerium concentrate for blocking layer screen printing", "91.5", "kg"),
        ("samarium europium gadolinium concentrate for blocking layer screen printing", "22.9", "kg"),
        ("aluminium oxide for electrolyte and H2 electrode", "6.4", "kg"),
        ("boric oxide for electrolyte and H2 electrode", "6.4", "kg"),
        ("barium oxide for electrolyte and H2 electrode", "6.4", "kg"),
        ("silicone product for electrolyte and H2 electrode", "6.4", "kg"),
        ("nickel for electrolyte and H2 electrode", "136.6", "kg"),
    ],
}


def render_table(title: str, rows: list[tuple[str, str, str]]) -> bytes:
    width = 1100
    row_h = 30
    height = 90 + row_h * (len(rows) + 1)
    pdf = fitz.open()
    page = pdf.new_page(width=width, height=height)
    page.insert_text((28, 32), title, fontsize=16)
    y = 66
    page.insert_text((28, y), "Foreground material / component", fontsize=11)
    page.insert_text((760, y), "Amount", fontsize=11)
    page.insert_text((900, y), "Unit", fontsize=11)
    y += row_h
    for name, amount, unit in rows:
        page.insert_text((28, y), name, fontsize=10)
        page.insert_text((760, y), amount, fontsize=10)
        page.insert_text((900, y), unit, fontsize=10)
        y += row_h
    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
    return pix.tobytes("png")


def main() -> None:
    out = Path(__file__).with_name("source_visual_inventory.docx")
    doc = Document()
    doc.add_paragraph("Gerloff 2021 supplementary inventory figures - normalized visual CI fixture")
    for title, rows in TABLES.items():
        doc.add_paragraph(f"Figure: Values of the 1 MW {title.split()[0]} system inventory")
        doc.add_picture(io.BytesIO(render_table(title, rows)))
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
