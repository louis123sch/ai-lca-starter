from __future__ import annotations

import argparse
import io
import json
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

import ai_lca.documents as document_module
from ai_lca.benchmark import evaluate_extraction, report_to_dict
from ai_lca.llm import extract_inventory_from_documents

TARGET_DOI = "10.1007/s11367-025-02462-7"
TARGET_SLUG = "eb1b5b25c4827b8d"


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _clean(element: ET.Element | None) -> str:
    if element is None:
        return ""
    return " ".join("".join(element.itertext()).split())


def _cell_lines(cell: ET.Element) -> list[str]:
    lines = []
    for child in list(cell):
        if _local(child.tag) == "p":
            text = _clean(child)
            if text:
                lines.append(text)
    if not lines:
        text = _clean(cell)
        if text:
            lines.append(text)
    return lines


def _article(xml_bytes: bytes) -> ET.Element:
    root = ET.fromstring(xml_bytes)
    if _local(root.tag) == "article":
        return root
    return next(element for element in root.iter() if _local(element.tag) == "article")


def _direct_child(element: ET.Element, tag: str) -> ET.Element | None:
    return next((child for child in list(element) if _local(child.tag) == tag), None)


def _build_document(xml_bytes: bytes) -> bytes:
    """Build a faithful machine-readable DOCX from the frozen JATS article.

    Paragraphs and table-cell paragraph boundaries are retained. This makes the test use
    the same DOCX ingestion path as the interactive app while keeping the frozen paper
    source fixed and reproducible.
    """
    article = _article(xml_bytes)
    doc = Document()
    title = next((e for e in article.iter() if _local(e.tag) == "article-title"), None)
    doc.add_heading(_clean(title), level=1)
    abstract = next((e for e in article.iter() if _local(e.tag) == "abstract"), None)
    if abstract is not None:
        doc.add_heading("Abstract", level=2)
        doc.add_paragraph(_clean(abstract))

    for section in article.iter():
        if _local(section.tag) != "sec":
            continue
        title_element = _direct_child(section, "title")
        title_text = _clean(title_element)
        if title_text:
            doc.add_heading(title_text, level=2)
        for child in list(section):
            if _local(child.tag) == "p":
                text = _clean(child)
                if text:
                    doc.add_paragraph(text)
            elif _local(child.tag) == "list":
                for item in child.iter():
                    if _local(item.tag) == "list-item":
                        text = _clean(item)
                        if text:
                            doc.add_paragraph(text, style="List Bullet")

    for wrap in (e for e in article.iter() if _local(e.tag) == "table-wrap"):
        label = _clean(_direct_child(wrap, "label"))
        caption = _clean(_direct_child(wrap, "caption"))
        doc.add_paragraph(f"{label}: {caption}".strip(": "))
        source_rows = [e for e in wrap.iter() if _local(e.tag) == "tr"]
        if not source_rows:
            continue
        column_count = max(
            len([c for c in list(row) if _local(c.tag) in {"td", "th"}])
            for row in source_rows
        )
        table = doc.add_table(rows=0, cols=column_count)
        for source_row in source_rows:
            source_cells = [c for c in list(source_row) if _local(c.tag) in {"td", "th"}]
            target_row = table.add_row()
            for column_index, target_cell in enumerate(target_row.cells):
                target_cell.text = ""
                if column_index >= len(source_cells):
                    continue
                lines = _cell_lines(source_cells[column_index])
                if not lines:
                    continue
                target_cell.paragraphs[0].add_run(lines[0])
                for line in lines[1:]:
                    target_cell.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _candidate_extract_docx_text(docx_bytes: bytes) -> str:
    """Candidate ingestion: preserve indexed paragraph structure inside table cells."""
    if not docx_bytes:
        raise ValueError("DOCX is empty")
    document = Document(io.BytesIO(docx_bytes))
    chunks: list[str] = []
    paragraph_no = 0
    table_no = 0

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            block = Paragraph(child, document)
            text = block.text.strip()
            if text:
                paragraph_no += 1
                chunks.append(f"[PARAGRAPH {paragraph_no}]\n{text}")
        elif isinstance(child, CT_Tbl):
            block = Table(child, document)
            table_no += 1
            rows = []
            for row_no, row in enumerate(block.rows, 1):
                rendered_cells = []
                for cell in row.cells:
                    lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                    if len(lines) <= 1:
                        rendered_cells.append(lines[0] if lines else "")
                    else:
                        rendered_cells.append(
                            " ⏎ ".join(f"[{index}] {line}" for index, line in enumerate(lines, 1))
                        )
                if any(rendered_cells):
                    rows.append(f"[ROW {row_no}] " + "\t".join(rendered_cells))
            if rows:
                chunks.append(f"[TABLE {table_no}]\n" + "\n".join(rows))

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError("No readable text was found in the DOCX")
    return combined


def _number(value: str) -> float | None:
    value = value.strip()
    if value in {"", "-", "–", "—"}:
        return None
    return float(value.replace(",", ""))


def _gold_from_source(xml_bytes: bytes) -> dict:
    article = _article(xml_bytes)
    inventory_table = None
    for wrap in (e for e in article.iter() if _local(e.tag) == "table-wrap"):
        label = _clean(_direct_child(wrap, "label"))
        caption = _clean(_direct_child(wrap, "caption"))
        if label.casefold() == "table 2" or "inventory of forklift" in caption.casefold():
            inventory_table = wrap
            break
    if inventory_table is None:
        raise RuntimeError("Could not locate the paper's foreground inventory table")

    flows: list[dict] = []
    source_rows = [e for e in inventory_table.iter() if _local(e.tag) == "tr"]
    for source_row in source_rows[1:]:
        cells = [c for c in list(source_row) if _local(c.tag) in {"td", "th"}]
        if len(cells) < 4:
            continue
        columns = [_cell_lines(cell) for cell in cells]
        resources, units, forklift_values, trailer_values = columns[:4]
        if not resources:
            continue
        if len(resources) > 1:
            for index, resource in enumerate(resources):
                # The first item is a reported aggregate total. Its component gases are the
                # actual independently quantified inventory exchanges, so do not double count it.
                if index == 0 and "total" in resource.casefold():
                    continue
                unit = units[index] if index < len(units) else (units[-1] if units else None)
                forklift_amount = _number(forklift_values[index]) if index < len(forklift_values) else None
                trailer_amount = _number(trailer_values[index]) if index < len(trailer_values) else None
                if forklift_amount is not None:
                    flows.append({"process_key": "forklift", "name": resource, "amount": forklift_amount, "unit": unit, "direction": "input"})
                if trailer_amount is not None:
                    flows.append({"process_key": "trailer", "name": resource, "amount": trailer_amount, "unit": unit, "direction": "input"})
        else:
            resource = resources[0]
            unit = units[0] if units else None
            forklift_amount = _number(forklift_values[0]) if forklift_values else None
            trailer_amount = _number(trailer_values[0]) if trailer_values else None
            if forklift_amount is not None:
                flows.append({"process_key": "forklift", "name": resource, "amount": forklift_amount, "unit": unit, "direction": "input"})
            if trailer_amount is not None:
                flows.append({"process_key": "trailer", "name": resource, "amount": trailer_amount, "unit": unit, "direction": "input"})

    # The paper also explicitly assesses a production-stage electric-forklift alternative.
    flows.append({"process_key": "electric", "name": "lithium-ion battery", "amount": None, "unit": None, "direction": "input"})
    return {
        "benchmark_id": "random44-heavy-machinery-document-read",
        "processes": [
            {"key": "forklift", "name": "Forklift manufacturing", "aliases": ["forklift manufacturing", "diesel forklift", "forklift"]},
            {"key": "trailer", "name": "Semi-trailer manufacturing", "aliases": ["semi-trailer manufacturing", "semi-trailer", "trailer manufacturing"]},
            {"key": "electric", "name": "Electric forklift production scenario", "aliases": ["electric forklift", "battery-equipped electric forklift", "electric forklift production"]},
        ],
        "flows": flows,
        "context": {"system_boundary": "cradle-to-gate", "reference_geography": "Türkiye"},
        "forbidden_foreground_name_terms": ["gwp", "carbon footprint", "co2 eq per"],
    }


def _run(docx_bytes: bytes, expected: dict, output_dir: Path, label: str, *, candidate: bool, model: str) -> dict:
    original = document_module.extract_docx_text
    if candidate:
        document_module.extract_docx_text = _candidate_extract_docx_text
    try:
        extraction = extract_inventory_from_documents(
            [("selected_paper.docx", docx_bytes)],
            model=model,
            max_visual_assets=0,
        )
    finally:
        document_module.extract_docx_text = original

    report = evaluate_extraction(extraction, expected)
    (output_dir / f"{label}_extraction.json").write_text(extraction.model_dump_json(indent=2) + "\n")
    payload = report_to_dict(report)
    (output_dir / f"{label}_report.json").write_text(json.dumps(payload, indent=2) + "\n")
    return payload


def _delta(before: dict, after: dict) -> dict:
    keys = ["overall_score", "process_recall", "process_precision", "flow_recall", "flow_precision", "amount_accuracy", "unit_accuracy", "direction_accuracy"]
    return {key: float(after[key]) - float(before[key]) for key in keys}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--artifact-root", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, default=Path("random44_reader_iteration"))
    parser.add_argument("--model", default="gpt-5-mini")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    source = args.artifact_root / "corpus" / TARGET_SLUG / "source" / "article.xml"
    if not source.exists():
        raise FileNotFoundError(f"Frozen target source not found: {source}")
    xml_bytes = source.read_bytes()
    document_bytes = _build_document(xml_bytes)
    expected = _gold_from_source(xml_bytes)
    (args.output_dir / "expected.json").write_text(json.dumps(expected, indent=2) + "\n")
    (args.output_dir / "selected_paper.docx").write_bytes(document_bytes)

    baseline = _run(document_bytes, expected, args.output_dir, "baseline", candidate=False, model=args.model)
    candidate = _run(document_bytes, expected, args.output_dir, "candidate", candidate=True, model=args.model)
    delta = _delta(baseline, candidate)
    accepted = (
        candidate["flow_recall"] > baseline["flow_recall"]
        and candidate["flow_precision"] >= baseline["flow_precision"] - 0.02
        and candidate["process_recall"] >= baseline["process_recall"]
        and candidate["process_precision"] >= baseline["process_precision"] - 0.05
        and candidate["amount_accuracy"] >= baseline["amount_accuracy"] - 0.05
    )
    result = {
        "doi": TARGET_DOI,
        "title": "Cradle-to-gate life cycle assessment of heavy machinery manufacturing: a case study in Türkiye",
        "hypothesis": "Preserving indexed intra-cell subrows in DOCX tables improves recovery of compound inventory rows without materially reducing precision.",
        "expected_flow_count": len(expected["flows"]),
        "baseline": baseline,
        "candidate": candidate,
        "delta": delta,
        "accepted": accepted,
        "model_calls_expected": 4,
    }
    (args.output_dir / "comparison.json").write_text(json.dumps(result, indent=2) + "\n")
    print(json.dumps(result, indent=2))
    if not accepted:
        raise SystemExit(3)


if __name__ == "__main__":
    main()
