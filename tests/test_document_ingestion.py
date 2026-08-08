import io

from docx import Document

from ai_lca.documents import extract_document_text, combine_document_texts


def test_extract_docx_text_with_source_marker():
    doc = Document()
    doc.add_paragraph("Plant lifetime is 20 years.")
    buffer = io.BytesIO()
    doc.save(buffer)

    text = extract_document_text(buffer.getvalue(), "example.docx")

    assert "[SOURCE example.docx]" in text
    assert "Plant lifetime is 20 years." in text


def test_extract_docx_tables():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Material"
    table.cell(0, 1).text = "Mass"
    table.cell(1, 0).text = "Steel"
    table.cell(1, 1).text = "100 kg"
    buffer = io.BytesIO()
    doc.save(buffer)

    text = extract_document_text(buffer.getvalue(), "inventory.docx")

    assert "[TABLE 1]" in text
    assert "Steel | 100 kg" in text


def test_combine_document_texts_keeps_sources_separate():
    combined = combine_document_texts(
        [
            ("one.docx", "[SOURCE one.docx]\nFirst document"),
            ("two.pdf", "[SOURCE two.pdf]\n[PAGE 1]\nSecond document"),
        ]
    )

    assert "[SOURCE one.docx]" in combined
    assert "[SOURCE two.pdf]" in combined
    assert "\n\n===== DOCUMENT BOUNDARY =====\n\n" in combined
