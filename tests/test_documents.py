import base64
import io

import fitz
from docx import Document

from ai_lca.documents import (
    combine_document_evidence,
    extract_document_text,
    extract_docx_text,
    extract_docx_visual_assets,
    extract_pdf_text,
)


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def test_pdf_extraction_keeps_page_markers():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Nickel input: 2.5 kg")
    pdf_bytes = doc.tobytes()
    text = extract_pdf_text(pdf_bytes)
    assert "[PAGE 1]" in text
    assert "Nickel input" in text


def test_docx_extraction_keeps_paragraph_and_table_markers():
    document = Document()
    document.add_paragraph("Natural gas input: 1.2 kg")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Electricity"
    table.cell(0, 1).text = "5 kWh"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_docx_text(buffer.getvalue())
    assert "[PARAGRAPH 1]" in text
    assert "Natural gas input" in text
    assert "[TABLE 1]" in text
    assert "Electricity\t5 kWh" in text


def test_docx_visual_extraction_keeps_caption_context():
    document = Document()
    document.add_paragraph("Inventory figure follows")
    document.add_picture(io.BytesIO(_ONE_PIXEL_PNG))
    document.add_paragraph("Figure 1: LCI values and material inputs")
    buffer = io.BytesIO()
    document.save(buffer)

    assets, warnings = extract_docx_visual_assets(buffer.getvalue(), filename="supplement.docx")
    assert not warnings
    assert len(assets) == 1
    assert assets[0].document == "supplement.docx"
    assert assets[0].mime_type == "image/png"
    assert "LCI values" in (assets[0].context or "")


def test_document_dispatch_accepts_docx():
    document = Document()
    document.add_paragraph("Hydrogen production")
    buffer = io.BytesIO()
    document.save(buffer)
    assert "Hydrogen production" in extract_document_text(buffer.getvalue(), "paper.docx")


def test_document_dispatch_accepts_utf8_text_fixture():
    text = "[PAGE 11]\nLCI data of SMR-based hydrogen production"
    assert extract_document_text(text.encode("utf-8"), "paper_fixture.txt") == text


def test_combined_documents_keep_document_markers():
    from ai_lca.documents import combine_document_texts

    document = Document()
    document.add_paragraph("Supplementary inventory")
    buffer = io.BytesIO()
    document.save(buffer)

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Main paper inventory")
    pdf_bytes = pdf.tobytes()

    text = combine_document_texts([
        ("paper.pdf", pdf_bytes),
        ("supplement.docx", buffer.getvalue()),
    ])
    assert "[DOCUMENT: paper.pdf]" in text
    assert "[DOCUMENT: supplement.docx]" in text
    assert "Main paper inventory" in text
    assert "Supplementary inventory" in text


def test_combined_evidence_returns_text_and_visual_assets():
    document = Document()
    document.add_paragraph("Supplementary inventory")
    document.add_picture(io.BytesIO(_ONE_PIXEL_PNG))
    document.add_paragraph("Figure 2: material inventory table")
    buffer = io.BytesIO()
    document.save(buffer)

    text, assets, warnings = combine_document_evidence([("supplement.docx", buffer.getvalue())])
    assert "[DOCUMENT: supplement.docx]" in text
    assert len(assets) == 1
    assert not warnings
