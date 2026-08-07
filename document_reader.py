from __future__ import annotations

import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def extract_pdf_text(pdf_bytes: bytes, max_pages: int | None = None, *, source_name: str | None = None) -> str:
    """Extract page-tagged text from a text-readable PDF locally."""
    if not pdf_bytes:
        raise ValueError("PDF is empty")

    doc = fitz.open(stream=io.BytesIO(pdf_bytes), filetype="pdf")
    chunks: list[str] = []
    if source_name:
        chunks.append(f"[SOURCE {source_name}]")

    n_pages = len(doc) if max_pages is None else min(len(doc), max_pages)
    for idx in range(n_pages):
        page = doc[idx]
        text = page.get_text("text").strip()
        chunks.append(f"[PAGE {idx + 1}]\n{text}")

    combined = "\n\n".join(chunks).strip()
    if not combined or combined == f"[SOURCE {source_name}]":
        raise ValueError(
            "No machine-readable text was found. This PDF may be scanned/image-only; "
            "use pasted text for now or add multimodal/OCR ingestion later."
        )
    return combined


def extract_docx_text(docx_bytes: bytes, *, source_name: str | None = None) -> str:
    """Extract paragraphs and tables from a .docx file while retaining source provenance."""
    if not docx_bytes:
        raise ValueError("Word document is empty")

    document = Document(io.BytesIO(docx_bytes))
    chunks: list[str] = []
    if source_name:
        chunks.append(f"[SOURCE {source_name}]")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            chunks.append(f"[TABLE {table_index}]\n" + "\n".join(rows))

    combined = "\n\n".join(chunks).strip()
    if not combined or combined == f"[SOURCE {source_name}]":
        raise ValueError("No readable text was found in the Word document.")
    return combined


def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch document extraction based on filename extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_bytes, source_name=filename)
    if suffix == ".docx":
        return extract_docx_text(file_bytes, source_name=filename)
    raise ValueError(f"Unsupported document type '{suffix}'. Supported types are PDF and DOCX.")
